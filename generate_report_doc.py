import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt

def generate_docx_from_html(html, output_file):
    soup = BeautifulSoup(html, 'html.parser')
    doc = Document()

    for element in soup.find_all(recursive=False):
        if element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'h3':
            doc.add_heading(element.text, level=3)
        elif element.name == 'p':
            doc.add_paragraph(element.text)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Bullet')
        elif element.name == 'img':
            doc.add_paragraph(f"[Image placeholder: {element['alt']}]")

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(12)

    doc.save(output_file)
    print(f"Word document generated successfully as '{output_file}'.")

if __name__ == "__main__":
    md_file = "plant_disease_recognition_report.md"
    docx_file = "plant_disease_recognition_report.docx"

    with open(md_file, 'r', encoding='utf-8') as file:
        md_content = file.read()

    html = markdown.markdown(md_content)
    generate_docx_from_html(html, docx_file)
